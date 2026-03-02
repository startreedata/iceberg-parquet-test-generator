"""One-shot script to generate architecture-overview.docx for team sharing."""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)


def build() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    title = doc.add_heading("Parquet-to-Pinot Type Testing", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph(
        "Architecture overview for Phase 1 functional testing of Parquet data type "
        "support in StarTree Pinot. Target: raw S3 catalog table type (Parquet files "
        "in S3, no Iceberg catalog metadata)."
    )

    # ---- 1. Purpose ----
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "Validate that every Parquet data type round-trips correctly through the "
        "Parquet → Pinot ingestion pipeline. The tool generates:"
    )
    doc.add_paragraph("Parquet test datasets with all physical, logical, and complex types", style="List Bullet")
    doc.add_paragraph("Pinot schema JSON matching the generated dataset", style="List Bullet")
    doc.add_paragraph("Pinot table config JSON ready for table creation", style="List Bullet")
    doc.add_paragraph(
        "The workflow is: generate dataset → create Pinot table with generated schema/config → "
        "ingest Parquet file → query and verify all data types are handled correctly."
    )

    # ---- 2. Architecture ----
    doc.add_heading("2. Architecture", level=1)
    doc.add_paragraph(
        "Config-driven Python tool using PyArrow and Pydantic. JSON configs specify "
        "Parquet-native types (physical + logical), not Pinot types."
    )
    doc.add_heading("Data Flow", level=2)
    doc.add_paragraph(
        "JSON Config → [config_model.py validates] → [schema_builder.py builds PyArrow schema] → "
        "[data_generator.py generates random typed data] → [parquet_writer.py writes .parquet] → "
        "[pinot_config_generator.py maps types to Pinot schema + table config]"
    )

    doc.add_heading("Key Design Decisions", level=2)
    doc.add_paragraph("Configs use Parquet-native types (physicalType: INT32, logicalType: DATE), not Pinot types", style="List Bullet")
    doc.add_paragraph("Seeded PRNG — same config + same seed = byte-identical output every time", style="List Bullet")
    doc.add_paragraph("Complex types (STRUCT, LIST, MAP) support arbitrary nesting", style="List Bullet")
    doc.add_paragraph("Null injection via per-column nullRatio; REQUIRED columns never get nulls", style="List Bullet")
    doc.add_paragraph("fixedValues field enables sentinel-collision testing (e.g. Integer.MIN_VALUE as real data)", style="List Bullet")

    # ---- 3. Parquet Type System ----
    doc.add_heading("3. Parquet Type System", level=1)
    doc.add_paragraph(
        "Parquet uses a layered type system: 8 physical types (storage encoding), "
        "logical types (semantic annotations), and complex types (group-based: STRUCT, LIST, MAP)."
    )

    doc.add_heading("3.1 Physical Types (8 total)", level=2)
    _add_table(doc,
        ["Physical Type", "Size", "Pinot Mapping"],
        [
            ["BOOLEAN", "1 bit", "BOOLEAN"],
            ["INT32", "32 bits", "INT"],
            ["INT64", "64 bits", "LONG"],
            ["FLOAT", "32 bits", "FLOAT"],
            ["DOUBLE", "64 bits", "DOUBLE"],
            ["BINARY", "variable", "STRING"],
            ["FIXED_LEN_BYTE_ARRAY", "fixed N bytes", "STRING"],
            ["INT96", "96 bits (deprecated)", "TIMESTAMP"],
        ],
    )

    doc.add_heading("3.2 Logical Types — Integers", level=2)
    _add_table(doc,
        ["Logical Type", "Physical", "Signed", "Pinot Mapping"],
        [
            ["INT(8, true)", "INT32", "yes", "INT"],
            ["INT(16, true)", "INT32", "yes", "INT"],
            ["INT(32, true)", "INT32", "yes", "INT"],
            ["INT(64, true)", "INT64", "yes", "LONG"],
            ["INT(8, false)", "INT32", "no", "INT"],
            ["INT(16, false)", "INT32", "no", "INT"],
            ["INT(32, false)", "INT32", "no", "LONG"],
            ["INT(64, false)", "INT64", "no", "LONG"],
        ],
    )
    doc.add_paragraph("Note: UINT32 → Pinot LONG (not INT) to avoid overflow since Pinot INT is signed 32-bit.")

    doc.add_heading("3.3 Logical Types — String & Binary", level=2)
    _add_table(doc,
        ["Logical Type", "Physical", "Pinot Mapping"],
        [
            ["STRING", "BINARY", "STRING"],
            ["ENUM", "BINARY", "STRING"],
            ["JSON", "BINARY", "STRING"],
            ["BSON", "BINARY", "BYTES"],
            ["UUID", "FIXED_LEN_BYTE_ARRAY(16)", "STRING"],
        ],
    )

    doc.add_heading("3.4 Logical Types — Temporal", level=2)
    _add_table(doc,
        ["Logical Type", "Physical", "Unit", "Pinot Mapping"],
        [
            ["DATE", "INT32", "days since epoch", "LONG"],
            ["TIME(MILLIS)", "INT32", "ms since midnight", "INT"],
            ["TIME(MICROS)", "INT64", "μs since midnight", "LONG"],
            ["TIME(NANOS)", "INT64", "ns since midnight", "LONG"],
            ["TIMESTAMP(MILLIS)", "INT64", "ms since epoch", "TIMESTAMP"],
            ["TIMESTAMP(MICROS)", "INT64", "μs since epoch", "TIMESTAMP"],
            ["TIMESTAMP(NANOS)", "INT64", "ns since epoch", "TIMESTAMP"],
        ],
    )
    doc.add_paragraph("Pinot TIMESTAMP stores milliseconds internally. Microsecond/nanosecond timestamps lose precision.")

    doc.add_heading("3.5 Logical Types — Decimal", level=2)
    _add_table(doc,
        ["Physical Backend", "Precision Range", "Pinot Mapping"],
        [
            ["INT32", "1–9", "STRING"],
            ["INT64", "1–18", "STRING"],
            ["FIXED_LEN_BYTE_ARRAY", "1–38", "STRING"],
            ["BINARY", "1–38", "STRING"],
        ],
    )
    doc.add_paragraph("ParquetToPinotTypeMapper maps DECIMAL → STRING. IcebergSchemaConverter maps DECIMAL → BIG_DECIMAL.")

    # ---- 4. Complex Types in Pinot ----
    doc.add_heading("4. Complex Type Support in Pinot", level=1)
    doc.add_paragraph(
        "Pinot does not natively store nested/hierarchical data. Complex Parquet types "
        "go through a lossy conversion. Via ParquetToPinotTypeMapper (the raw S3 path), "
        "ALL complex types become STRING (JSON serialized)."
    )

    doc.add_heading("4.1 ParquetToPinotTypeMapper (Raw S3 / Default)", level=2)
    _add_table(doc,
        ["Parquet Type", "Pinot Type", "Notes"],
        [
            ["STRUCT", "STRING (JSON)", "Always serialized to JSON"],
            ["LIST (any element type)", "STRING (JSON)", "JSON serialized, not MV column"],
            ["MAP (any key/value type)", "STRING (JSON)", "JSON serialized"],
            ["Nested complex", "STRING (JSON)", "All nesting collapses to JSON blob"],
        ],
    )

    doc.add_heading("4.2 IcebergSchemaConverter (Iceberg Pipeline)", level=2)
    doc.add_paragraph("When using Iceberg catalog metadata, the mapping is more granular:")
    _add_table(doc,
        ["Parquet Type", "Pinot Type", "Notes"],
        [
            ["LIST of primitives", "Multi-value DimensionFieldSpec", "e.g. list<string> → MV STRING"],
            ["LIST of complex", "JSON", "Serialized"],
            ["MAP<prim, prim>", "MAP (ComplexFieldSpec)", "Native key/value child fields"],
            ["MAP with complex values", "JSON", "Serialized"],
            ["STRUCT", "JSON", "Always serialized"],
        ],
    )

    doc.add_heading("4.3 Known Gaps", level=2)
    doc.add_paragraph("ParquetSequentialColumnReaderFactory does not handle multi-value or struct columns yet (Arrow-based reader).", style="List Bullet")
    doc.add_paragraph("STRUCT always becomes JSON — queryable via JSON indexing but loses the typed schema.", style="List Bullet")
    doc.add_paragraph("Deeply nested types (list of structs, map of maps) all collapse to opaque JSON.", style="List Bullet")
    doc.add_paragraph("MulivalueFieldConverter does not check sentinel nulls — only primitive converters do.", style="List Bullet")

    # ---- 5. Complete Mapping Table ----
    doc.add_heading("5. Complete Parquet → Pinot Mapping (ParquetToPinotTypeMapper)", level=1)
    _add_table(doc,
        ["Parquet Type", "Pinot Type", "Notes"],
        [
            ["BOOLEAN", "BOOLEAN", ""],
            ["INT32", "INT", ""],
            ["INT32 + DATE", "LONG", "Days since epoch"],
            ["INT32 + TIME(MILLIS)", "INT", "ms since midnight"],
            ["INT32 + DECIMAL(P,S)", "STRING", "P ≤ 9"],
            ["INT32 + INT(8/16, signed)", "INT", ""],
            ["INT32 + UINT32", "LONG", "Widened to avoid overflow"],
            ["INT64", "LONG", ""],
            ["INT64 + TIMESTAMP(MILLIS)", "TIMESTAMP", ""],
            ["INT64 + TIMESTAMP(MICROS)", "TIMESTAMP", ""],
            ["INT64 + TIMESTAMP(NANOS)", "TIMESTAMP", "Loses precision"],
            ["INT64 + TIME(MICROS)", "LONG", ""],
            ["INT64 + TIME(NANOS)", "LONG", ""],
            ["INT64 + DECIMAL(P,S)", "STRING", "P ≤ 18"],
            ["FLOAT", "FLOAT", ""],
            ["DOUBLE", "DOUBLE", ""],
            ["BINARY + STRING", "STRING", ""],
            ["BINARY + ENUM", "STRING", ""],
            ["BINARY + JSON", "STRING", ""],
            ["BINARY + BSON", "BYTES", ""],
            ["BINARY (no annotation)", "STRING", ""],
            ["FIXED + UUID", "STRING", "16 bytes"],
            ["FIXED + DECIMAL(P,S)", "STRING", "P ≤ 38"],
            ["FIXED (no annotation)", "STRING", ""],
            ["INT96", "TIMESTAMP", "Deprecated, no stats"],
            ["STRUCT", "STRING (JSON)", ""],
            ["LIST", "STRING (JSON)", ""],
            ["MAP", "STRING (JSON)", ""],
        ],
    )

    # ---- 6. Null Handling ----
    doc.add_heading("6. Null Handling", level=1)

    doc.add_heading("6.1 Parquet Side", level=2)
    doc.add_paragraph(
        "Parquet uses definition levels to encode nulls. REQUIRED columns have no definition "
        "level (saves space), OPTIONAL columns track presence per row. Nested types (LIST, MAP, "
        "STRUCT) have multi-level definitions distinguishing null-container vs null-element."
    )

    doc.add_heading("6.2 Pinot Side — Sentinel-Based (Legacy)", level=2)
    doc.add_paragraph(
        "When enableColumnBasedNullHandling is false, Pinot uses sentinel values:"
    )
    _add_table(doc,
        ["Pinot Type", "Sentinel Value", "Risk"],
        [
            ["INT", "Integer.MIN_VALUE", "Collision with real data"],
            ["LONG", "Long.MIN_VALUE", "Collision with real data"],
            ["FLOAT", "Float.NEGATIVE_INFINITY", "Rare but possible"],
            ["DOUBLE", "Double.NEGATIVE_INFINITY", "Rare but possible"],
            ["STRING", '"null" literal', 'String "null" indistinguishable from null'],
            ["BOOLEAN", "-1 (as int)", ""],
        ],
    )

    doc.add_heading("6.3 Pinot Side — Column-Based (Recommended)", level=2)
    doc.add_paragraph(
        "When enableColumnBasedNullHandling is true, Pinot uses a null bitmap per column. "
        "No sentinel collisions, proper IS NULL / IS NOT NULL semantics. This is what the "
        "generated table configs enable."
    )

    # ---- 7. Edge Cases ----
    doc.add_heading("7. Edge Cases & Known Limitations", level=1)

    items = [
        ("INT96 — no Parquet statistics",
         "INT96 columns lack column statistics. Cannot be used as timeColumnName in Pinot "
         "table config or IcebergWatcher fails with 'No time statistics available'. "
         "The generator treats INT96 as a dimension, not dateTime."),
        ("Timestamp precision mismatch",
         "Iceberg defaults to microsecond precision. Pinot TIMESTAMP stores milliseconds "
         "internally. Nanosecond timestamps lose precision on ingestion."),
        ("UTC vs local timestamps",
         "Pinot treats isAdjustedToUTC=true and false identically (both become TIMESTAMP), "
         "losing timezone semantics."),
        ("BINARY defaults to STRING",
         "BINARY and FIXED_LEN_BYTE_ARRAY without logical type annotation both default to "
         "STRING in Pinot, not BYTES. Only BSON maps to BYTES."),
        ("DECIMAL precision preservation",
         "ParquetToPinotTypeMapper maps DECIMAL → STRING to preserve precision. "
         "IcebergSchemaConverter maps DECIMAL → BIG_DECIMAL."),
        ("Sentinel-collision risk",
         "With legacy null handling, Integer.MIN_VALUE / Long.MIN_VALUE / "
         "Float.NEGATIVE_INFINITY in real data are incorrectly treated as null."),
    ]
    for title_text, body_text in items:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(body_text)

    # ---- 8. Test Configs ----
    doc.add_heading("8. Pre-Built Test Configs", level=1)
    _add_table(doc,
        ["Config File", "What It Tests"],
        [
            ["all-primitive-types.json", "All 8 physical types (OPTIONAL + REQUIRED variants)"],
            ["all-logical-types.json", "All logical type combinations (integers, temporal, decimal, string/binary)"],
            ["all-complex-types.json", "STRUCT, LIST, MAP, nested combinations (list-of-struct, map-of-struct, etc.)"],
            ["iceberg-full-coverage.json", "Types as Iceberg would produce them (microsecond timestamps, decimal on FIXED)"],
            ["null-scenarios.json", "8 null edge cases: sentinel collision, nullable LIST/MAP/STRUCT, schema mismatch"],
            ["golden-dataset-extended.json", "Superset golden schema with all type categories"],
        ],
    )

    # ---- 9. CLI Usage ----
    doc.add_heading("9. CLI Usage", level=1)

    doc.add_heading("Generate Parquet data", level=2)
    doc.add_paragraph("python -m generator.cli --config configs/all-primitive-types.json --output output/", style="No Spacing")

    doc.add_heading("Generate Pinot schema + table config only", level=2)
    doc.add_paragraph("python -m generator.cli --config configs/all-primitive-types.json --output output/ --pinot-config", style="No Spacing")

    doc.add_heading("Generate both Parquet data and Pinot configs", level=2)
    doc.add_paragraph("python -m generator.cli --config configs/all-primitive-types.json --output output/ --pinot-config --parquet", style="No Spacing")

    doc.add_heading("Use Iceberg mapping mode", level=2)
    doc.add_paragraph("python -m generator.cli --config configs/iceberg-full-coverage.json --output output/ --pinot-config --mapping-mode iceberg", style="No Spacing")

    doc.add_heading("Output Files", level=2)
    _add_table(doc,
        ["File", "Description"],
        [
            ["{name}.parquet", "Generated Parquet dataset"],
            ["{name}_pinot_schema.json", "Pinot schema (dimensionFieldSpecs, dateTimeFieldSpecs, complexFieldSpecs)"],
            ["{name}_pinot_table_config.json", "OFFLINE table config (null handling, JSON indexes, time column)"],
        ],
    )

    # ---- 10. Testing Workflow ----
    doc.add_heading("10. Testing Workflow", level=1)
    doc.add_paragraph("1. Pick a config (e.g. all-primitive-types.json)", style="List Number")
    doc.add_paragraph("2. Generate Parquet + Pinot configs:  python -m generator.cli -c configs/all-primitive-types.json -o output/ --pinot-config --parquet", style="List Number")
    doc.add_paragraph("3. Upload the Parquet file to S3", style="List Number")
    doc.add_paragraph("4. Create the Pinot table using the generated schema and table config JSONs", style="List Number")
    doc.add_paragraph("5. Trigger segment ingestion from S3", style="List Number")
    doc.add_paragraph("6. Query the table to verify all data types round-tripped correctly", style="List Number")
    doc.add_paragraph("7. Repeat for each config to cover all type categories", style="List Number")

    # ---- Save ----
    out_path = "output/Parquet-to-Pinot-Type-Testing-Architecture.docx"
    doc.save(out_path)
    print(f"Written to: {out_path}")


if __name__ == "__main__":
    build()
