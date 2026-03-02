"""One-shot script to generate the test plan as a .docx with Status columns."""
from __future__ import annotations

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def _set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for i, w in enumerate(widths_cm):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]],
               col_widths: list[float] | None = None, font_size: int = 8) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    if headers[c_idx] == "Status":
                        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    if col_widths:
        _set_col_widths(table, col_widths)


WIDTHS_6COL = [1.0, 2.5, 2.5, 1.8, 7.0, 4.5]
WIDTHS_7COL = [1.0, 2.5, 2.5, 1.8, 7.0, 4.5, 1.8]
STATUS = ""  # blank for user to fill in


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    section = doc.sections[0]
    section.orientation = 1  # landscape
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ---- Title ----
    title = doc.add_heading("Parquet Data Type Testing — Test Plan", level=0)
    doc.add_paragraph(
        "Phase 1 functional testing of all Parquet data types supported by Pinot "
        "via raw S3 catalog table type."
    )

    # ---- Workflow ----
    doc.add_heading("Workflow per Test", level=1)
    doc.add_paragraph("Upload .parquet file to S3", style="List Number")
    doc.add_paragraph("Create Pinot table using the generated _pinot_schema.json and _pinot_table_config.json", style="List Number")
    doc.add_paragraph("Trigger segment ingestion from S3", style="List Number")
    doc.add_paragraph("Run validation queries below", style="List Number")
    doc.add_paragraph("Mark Status column as PASS / FAIL / BLOCKED", style="List Number")

    # ==================================================================
    # TEST 1: Primitive Types
    # ==================================================================
    doc.add_heading("Test 1: Primitive Types", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("all-primitive-types.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("100,000")

    _add_table(doc,
        ["#", "Column", "Parquet Type", "Pinot Type", "Validation Query", "Pass Criteria", "Status"],
        [
            ["1.1", "col_boolean", "BOOLEAN", "BOOLEAN",
             "SELECT col_boolean, COUNT(*) FROM all_primitive_types GROUP BY col_boolean",
             "true/false present, null ~10%", STATUS],
            ["1.2", "col_int32", "INT32", "INT",
             "SELECT MIN(col_int32), MAX(col_int32), COUNT(*) FROM all_primitive_types WHERE col_int32 IS NOT NULL",
             "Values in INT32 range, ~90k non-null", STATUS],
            ["1.3", "col_int64", "INT64", "LONG",
             "SELECT MIN(col_int64), MAX(col_int64), COUNT(*) FROM all_primitive_types WHERE col_int64 IS NOT NULL",
             "Values in INT64 range", STATUS],
            ["1.4", "col_float", "FLOAT", "FLOAT",
             "SELECT MIN(col_float), MAX(col_float), COUNT(*) FROM all_primitive_types WHERE col_float IS NOT NULL",
             "Float values present", STATUS],
            ["1.5", "col_double", "DOUBLE", "DOUBLE",
             "SELECT MIN(col_double), MAX(col_double), COUNT(*) FROM all_primitive_types WHERE col_double IS NOT NULL",
             "Double values present", STATUS],
            ["1.6", "col_binary", "BINARY", "STRING",
             "SELECT col_binary FROM all_primitive_types WHERE col_binary IS NOT NULL LIMIT 5",
             "Byte data readable as string", STATUS],
            ["1.7", "col_fixed_len", "FIXED(20)", "STRING",
             "SELECT col_fixed_len FROM all_primitive_types WHERE col_fixed_len IS NOT NULL LIMIT 5",
             "20-byte data as string", STATUS],
            ["1.8", "col_int96", "INT96 (deprecated)", "TIMESTAMP",
             "SELECT col_int96 FROM all_primitive_types WHERE col_int96 IS NOT NULL LIMIT 5",
             "Timestamp values readable", STATUS],
            ["1.9", "col_boolean_required", "BOOLEAN (REQ)", "BOOLEAN",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_boolean_required IS NULL",
             "Count = 0", STATUS],
            ["1.10", "col_int32_required", "INT32 (REQ)", "INT",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_int32_required IS NULL",
             "Count = 0", STATUS],
            ["1.11", "col_int64_required", "INT64 (REQ)", "LONG",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_int64_required IS NULL",
             "Count = 0", STATUS],
            ["1.12", "col_float_required", "FLOAT (REQ)", "FLOAT",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_float_required IS NULL",
             "Count = 0", STATUS],
            ["1.13", "col_double_required", "DOUBLE (REQ)", "DOUBLE",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_double_required IS NULL",
             "Count = 0", STATUS],
            ["1.14", "col_binary_required", "BINARY (REQ)", "STRING",
             "SELECT COUNT(*) FROM all_primitive_types WHERE col_binary_required IS NULL",
             "Count = 0", STATUS],
        ],
        col_widths=[1.0, 3.0, 2.5, 2.0, 8.0, 4.0, 1.8],
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Global check: ").bold = True
    p.add_run("SELECT COUNT(*) FROM all_primitive_types → 100,000")

    # ==================================================================
    # TEST 2: Logical Types
    # ==================================================================
    doc.add_heading("Test 2: Logical Types", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("all-logical-types.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("100,000")

    doc.add_heading("2A: Integer Subtypes", level=2)
    _add_table(doc,
        ["#", "Column", "Parquet Logical", "Pinot", "Validation Query", "Pass Criteria", "Status"],
        [
            ["2.1", "col_int8", "INT(8, signed)", "INT",
             "SELECT MIN(col_int8), MAX(col_int8) FROM all_logical_types WHERE col_int8 IS NOT NULL",
             "Range: -128 to 127", STATUS],
            ["2.2", "col_int16", "INT(16, signed)", "INT",
             "SELECT MIN(col_int16), MAX(col_int16) FROM all_logical_types WHERE col_int16 IS NOT NULL",
             "Range: -32768 to 32767", STATUS],
            ["2.3", "col_int32_logical", "INT(32, signed)", "INT",
             "SELECT MIN(col_int32_logical), MAX(col_int32_logical) FROM all_logical_types WHERE col_int32_logical IS NOT NULL",
             "INT32 range", STATUS],
            ["2.4", "col_int64_logical", "INT(64, signed)", "LONG",
             "SELECT MIN(col_int64_logical), MAX(col_int64_logical) FROM all_logical_types WHERE col_int64_logical IS NOT NULL",
             "INT64 range", STATUS],
            ["2.5", "col_uint8", "INT(8, unsigned)", "INT",
             "SELECT MIN(col_uint8), MAX(col_uint8) FROM all_logical_types WHERE col_uint8 IS NOT NULL",
             "Range: 0 to 255", STATUS],
            ["2.6", "col_uint16", "INT(16, unsigned)", "INT",
             "SELECT MIN(col_uint16), MAX(col_uint16) FROM all_logical_types WHERE col_uint16 IS NOT NULL",
             "Range: 0 to 65535", STATUS],
            ["2.7", "col_uint32", "INT(32, unsigned)", "LONG",
             "SELECT MIN(col_uint32), MAX(col_uint32) FROM all_logical_types WHERE col_uint32 IS NOT NULL",
             "Range: 0 to 4294967295", STATUS],
            ["2.8", "col_uint64", "INT(64, unsigned)", "LONG",
             "SELECT MIN(col_uint64), MAX(col_uint64) FROM all_logical_types WHERE col_uint64 IS NOT NULL",
             "Non-negative values", STATUS],
        ],
        col_widths=[1.0, 2.5, 2.5, 1.5, 8.5, 3.5, 1.8],
    )

    doc.add_heading("2B: String and Binary Types", level=2)
    _add_table(doc,
        ["#", "Column", "Parquet Logical", "Pinot", "Validation Query", "Pass Criteria", "Status"],
        [
            ["2.9", "col_string", "STRING", "STRING",
             "SELECT col_string FROM all_logical_types WHERE col_string IS NOT NULL LIMIT 5",
             "UTF-8 strings readable", STATUS],
            ["2.10", "col_enum", "ENUM", "STRING",
             "SELECT col_enum FROM all_logical_types WHERE col_enum IS NOT NULL LIMIT 5",
             "Enum values as strings", STATUS],
            ["2.11", "col_json", "JSON", "STRING",
             "SELECT col_json FROM all_logical_types WHERE col_json IS NOT NULL LIMIT 5",
             "Valid JSON strings", STATUS],
            ["2.12", "col_bson", "BSON", "BYTES",
             "SELECT col_bson FROM all_logical_types WHERE col_bson IS NOT NULL LIMIT 5",
             "Binary data present", STATUS],
            ["2.13", "col_uuid", "UUID", "STRING",
             "SELECT col_uuid FROM all_logical_types WHERE col_uuid IS NOT NULL LIMIT 5",
             "UUID-formatted strings", STATUS],
        ],
        col_widths=[1.0, 2.5, 2.5, 1.5, 8.5, 3.5, 1.8],
    )

    doc.add_heading("2C: Temporal Types", level=2)
    _add_table(doc,
        ["#", "Column", "Parquet Logical", "Pinot", "Validation Query", "Pass Criteria", "Status"],
        [
            ["2.14", "col_date", "DATE", "LONG",
             "SELECT MIN(col_date), MAX(col_date) FROM all_logical_types WHERE col_date IS NOT NULL",
             "Epoch day values, reasonable range", STATUS],
            ["2.15", "col_time_millis", "TIME(MILLIS)", "INT",
             "SELECT MIN(col_time_millis), MAX(col_time_millis) FROM all_logical_types WHERE col_time_millis IS NOT NULL",
             "0 to 86399999", STATUS],
            ["2.16", "col_time_micros", "TIME(MICROS)", "LONG",
             "SELECT MIN(col_time_micros), MAX(col_time_micros) FROM all_logical_types WHERE col_time_micros IS NOT NULL",
             "0 to 86399999999", STATUS],
            ["2.17", "col_time_nanos", "TIME(NANOS)", "LONG",
             "SELECT MIN(col_time_nanos), MAX(col_time_nanos) FROM all_logical_types WHERE col_time_nanos IS NOT NULL",
             "0 to 86399999999999", STATUS],
            ["2.18", "col_timestamp_millis_utc", "TS(MILLIS, UTC)", "TIMESTAMP",
             "SELECT col_timestamp_millis_utc FROM all_logical_types WHERE col_timestamp_millis_utc IS NOT NULL LIMIT 5",
             "Readable timestamps", STATUS],
            ["2.19", "col_timestamp_millis_local", "TS(MILLIS, local)", "TIMESTAMP",
             "SELECT col_timestamp_millis_local FROM all_logical_types WHERE col_timestamp_millis_local IS NOT NULL LIMIT 5",
             "Readable timestamps", STATUS],
            ["2.20", "col_timestamp_micros_utc", "TS(MICROS, UTC)", "TIMESTAMP",
             "SELECT col_timestamp_micros_utc FROM all_logical_types WHERE col_timestamp_micros_utc IS NOT NULL LIMIT 5",
             "Readable (micros→ms loss OK)", STATUS],
            ["2.21", "col_timestamp_micros_local", "TS(MICROS, local)", "TIMESTAMP",
             "SELECT col_timestamp_micros_local FROM all_logical_types WHERE col_timestamp_micros_local IS NOT NULL LIMIT 5",
             "Readable (micros→ms loss OK)", STATUS],
            ["2.22", "col_timestamp_nanos_utc", "TS(NANOS, UTC)", "TIMESTAMP",
             "SELECT col_timestamp_nanos_utc FROM all_logical_types WHERE col_timestamp_nanos_utc IS NOT NULL LIMIT 5",
             "Readable (nanos→ms loss OK)", STATUS],
            ["2.23", "col_timestamp_nanos_local", "TS(NANOS, local)", "TIMESTAMP",
             "SELECT col_timestamp_nanos_local FROM all_logical_types WHERE col_timestamp_nanos_local IS NOT NULL LIMIT 5",
             "Readable (nanos→ms loss OK)", STATUS],
        ],
        col_widths=[1.0, 3.5, 2.5, 1.8, 8.5, 3.5, 1.8],
    )

    doc.add_heading("2D: Decimal Types", level=2)
    _add_table(doc,
        ["#", "Column", "Parquet Type", "Pinot", "Validation Query", "Pass Criteria", "Status"],
        [
            ["2.24", "col_decimal_int32", "INT32+DECIMAL(9,2)", "BIG_DECIMAL",
             "SELECT col_decimal_int32 FROM all_logical_types WHERE col_decimal_int32 IS NOT NULL LIMIT 5",
             "Decimal values, scale=2", STATUS],
            ["2.25", "col_decimal_int64", "INT64+DECIMAL(18,4)", "BIG_DECIMAL",
             "SELECT col_decimal_int64 FROM all_logical_types WHERE col_decimal_int64 IS NOT NULL LIMIT 5",
             "Decimal values, scale=4", STATUS],
            ["2.26", "col_decimal_fixed", "FIXED+DECIMAL(28,6)", "BIG_DECIMAL",
             "SELECT col_decimal_fixed FROM all_logical_types WHERE col_decimal_fixed IS NOT NULL LIMIT 5",
             "Decimal values, scale=6", STATUS],
            ["2.27", "col_decimal_binary", "BIN+DECIMAL(38,10)", "BIG_DECIMAL",
             "SELECT col_decimal_binary FROM all_logical_types WHERE col_decimal_binary IS NOT NULL LIMIT 5",
             "Decimal values, scale=10", STATUS],
        ],
        col_widths=[1.0, 2.8, 3.0, 2.0, 8.5, 3.0, 1.8],
    )

    # ==================================================================
    # TEST 3: Complex Types
    # ==================================================================
    doc.add_heading("Test 3: Complex Types", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("all-complex-types.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("100,000    ")
    p.add_run("Note: ").bold = True
    p.add_run("All complex types map to JSON via ParquetToPinotTypeMapper")

    _add_table(doc,
        ["#", "Column", "Parquet Type", "Pinot", "Validation Query", "Pass Criteria", "Status"],
        [
            ["3.1", "flat_struct", "STRUCT{name,age,score}", "JSON",
             "SELECT flat_struct FROM all_complex_types WHERE flat_struct IS NOT NULL LIMIT 3",
             'Valid JSON w/ "name","age","score"', STATUS],
            ["3.2", "nested_struct", "STRUCT{id,address{...}}", "JSON",
             "SELECT nested_struct FROM all_complex_types WHERE nested_struct IS NOT NULL LIMIT 3",
             'Nested "address" object', STATUS],
            ["3.3", "string_list", "LIST<STRING>", "JSON",
             "SELECT string_list FROM all_complex_types WHERE string_list IS NOT NULL LIMIT 3",
             "JSON array of strings", STATUS],
            ["3.4", "int_list", "LIST<INT32>", "JSON",
             "SELECT int_list FROM all_complex_types WHERE int_list IS NOT NULL LIMIT 3",
             "JSON array of integers", STATUS],
            ["3.5", "list_of_structs", "LIST<STRUCT>", "JSON",
             "SELECT list_of_structs FROM all_complex_types WHERE list_of_structs IS NOT NULL LIMIT 3",
             "JSON array of objects", STATUS],
            ["3.6", "map_string_int", "MAP<STRING,INT>", "JSON",
             "SELECT map_string_int FROM all_complex_types WHERE map_string_int IS NOT NULL LIMIT 3",
             "JSON object, string keys, int vals", STATUS],
            ["3.7", "map_string_string", "MAP<STRING,STRING>", "JSON",
             "SELECT map_string_string FROM all_complex_types WHERE map_string_string IS NOT NULL LIMIT 3",
             "JSON object, string k/v", STATUS],
            ["3.8", "map_string_struct", "MAP<STRING,STRUCT>", "JSON",
             "SELECT map_string_struct FROM all_complex_types WHERE map_string_struct IS NOT NULL LIMIT 3",
             "JSON with nested struct values", STATUS],
            ["3.9", "struct_with_list", "STRUCT{label,tags:LIST}", "JSON",
             "SELECT struct_with_list FROM all_complex_types WHERE struct_with_list IS NOT NULL LIMIT 3",
             '"tags" as nested array', STATUS],
            ["3.10", "struct_with_map", "STRUCT{name,attrs:MAP}", "JSON",
             "SELECT struct_with_map FROM all_complex_types WHERE struct_with_map IS NOT NULL LIMIT 3",
             '"attrs" as nested object', STATUS],
        ],
        col_widths=[1.0, 2.8, 3.2, 1.2, 8.5, 3.5, 1.8],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("JSON index check: ").bold = True
    doc.add_paragraph(
        "SELECT JSON_EXTRACT_SCALAR(flat_struct, '$.name', 'STRING') AS name, "
        "JSON_EXTRACT_SCALAR(flat_struct, '$.age', 'INT') AS age "
        "FROM all_complex_types WHERE JSON_MATCH(flat_struct, '\"$.name\" IS NOT NULL') LIMIT 5",
        style="No Spacing",
    )

    # ==================================================================
    # TEST 4: Null Handling
    # ==================================================================
    doc.add_heading("Test 4: Null Handling", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("null-scenarios.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("100,000")

    doc.add_heading("4A: Basic Optional Primitives (~15% nulls)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.1", "scenario1_optional_int",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_int IS NULL",
             "~15,000 nulls", STATUS],
            ["4.2", "scenario1_optional_long",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_long IS NULL",
             "~15,000 nulls", STATUS],
            ["4.3", "scenario1_optional_float",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_float IS NULL",
             "~15,000 nulls", STATUS],
            ["4.4", "scenario1_optional_double",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_double IS NULL",
             "~15,000 nulls", STATUS],
            ["4.5", "scenario1_optional_bool",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_bool IS NULL",
             "~15,000 nulls", STATUS],
            ["4.6", "scenario1_optional_string",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario1_optional_string IS NULL",
             "~15,000 nulls", STATUS],
        ],
        col_widths=[1.0, 3.5, 10.0, 3.0, 1.8],
    )

    doc.add_heading("4B: Required Columns (Zero Nulls)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.7", "scenario2_required_int",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario2_required_int IS NULL",
             "0", STATUS],
            ["4.8", "scenario2_required_long",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario2_required_long IS NULL",
             "0", STATUS],
            ["4.9", "scenario2_required_string",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario2_required_string IS NULL",
             "0", STATUS],
        ],
        col_widths=[1.0, 3.5, 10.0, 3.0, 1.8],
    )

    doc.add_heading("4C: Sentinel Collision (~20% nulls, fixedValues include sentinels)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.10", "scenario3_sentinel_int",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario3_sentinel_int = -2147483648",
             "Rows with MIN_VALUE as real data (not treated as null)", STATUS],
            ["4.11", "scenario3_sentinel_long",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario3_sentinel_long = -9223372036854775808",
             "Rows with Long.MIN_VALUE as real data", STATUS],
            ["4.12", "scenario3_sentinel_string",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario3_sentinel_string = 'null'",
             'Literal "null" string preserved as data', STATUS],
        ],
        col_widths=[1.0, 3.5, 10.0, 5.0, 1.8],
    )

    doc.add_heading("4D: Nullable Complex Types (~20% nulls)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.13", "scenario4_nullable_list",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario4_nullable_list IS NULL",
             "~20,000 nulls", STATUS],
            ["4.14", "scenario4_nullable_list (non-null)",
             "SELECT scenario4_nullable_list FROM null_scenarios WHERE scenario4_nullable_list IS NOT NULL LIMIT 5",
             "Valid JSON arrays", STATUS],
            ["4.15", "scenario5_nullable_map",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario5_nullable_map IS NULL",
             "~20,000 nulls", STATUS],
            ["4.16", "scenario6_nullable_struct",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario6_nullable_struct IS NULL",
             "~20,000 nulls", STATUS],
        ],
        col_widths=[1.0, 4.5, 10.0, 3.0, 1.8],
    )

    doc.add_heading("4E: Nullable Temporal (~25% nulls)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.17", "scenario7_optional_date",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario7_optional_date IS NULL",
             "~25,000 nulls", STATUS],
            ["4.18", "scenario7_optional_timestamp",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario7_optional_timestamp IS NULL",
             "~25,000 nulls", STATUS],
        ],
        col_widths=[1.0, 4.0, 10.0, 3.0, 1.8],
    )

    doc.add_heading("4F: Required vs Optional Strings (~50% nulls for optional)", level=2)
    _add_table(doc,
        ["#", "Column", "Validation Query", "Pass Criteria", "Status"],
        [
            ["4.19", "scenario8_required_iceberg_string",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario8_required_iceberg_string IS NULL",
             "0", STATUS],
            ["4.20", "scenario8_optional_iceberg_string",
             "SELECT COUNT(*) FROM null_scenarios WHERE scenario8_optional_iceberg_string IS NULL",
             "~50,000 nulls", STATUS],
        ],
        col_widths=[1.0, 4.5, 10.0, 3.0, 1.8],
    )

    # ==================================================================
    # TEST 5: Iceberg Full Coverage
    # ==================================================================
    doc.add_heading("Test 5: Iceberg Full Coverage", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("iceberg-full-coverage.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("100,000")

    _add_table(doc,
        ["#", "Column", "Type Chain", "Validation", "Status"],
        [
            ["5.1", "iceberg_boolean", "BOOLEAN", "Standard boolean check", STATUS],
            ["5.2", "iceberg_int", "INT32+INT(32,true) → INT", "INT range", STATUS],
            ["5.3", "iceberg_long", "INT64+INT(64,true) → LONG", "LONG range", STATUS],
            ["5.4", "iceberg_float", "FLOAT", "Float values", STATUS],
            ["5.5", "iceberg_double", "DOUBLE", "Double values", STATUS],
            ["5.6", "iceberg_date", "INT32+DATE → LONG", "Epoch days, positive", STATUS],
            ["5.7", "iceberg_time", "INT64+TIME(MICROS) → LONG", "0 to 86399999999", STATUS],
            ["5.8", "iceberg_timestamp", "INT64+TS(MICROS,local) → TIMESTAMP", "Readable timestamps", STATUS],
            ["5.9", "iceberg_timestamptz", "INT64+TS(MICROS,UTC) → TIMESTAMP", "Readable timestamps", STATUS],
            ["5.10", "iceberg_string", "BINARY+STRING", "UTF-8 strings", STATUS],
            ["5.11", "iceberg_uuid", "FIXED(16)+UUID → STRING", "UUID-format strings", STATUS],
            ["5.12", "iceberg_binary", "BINARY → STRING", "Binary as string", STATUS],
            ["5.13", "iceberg_fixed_8", "FIXED(8) → STRING", "8-byte data", STATUS],
            ["5.14", "iceberg_decimal_9_2", "FIXED+DECIMAL(9,2) → BIG_DECIMAL", "Scale=2", STATUS],
            ["5.15", "iceberg_decimal_18_4", "FIXED+DECIMAL(18,4) → BIG_DECIMAL", "Scale=4", STATUS],
            ["5.16", "iceberg_decimal_38_10", "FIXED+DECIMAL(38,10) → BIG_DECIMAL", "Scale=10, high precision", STATUS],
            ["5.17", "iceberg_struct", "STRUCT → JSON", "Valid JSON object", STATUS],
            ["5.18", "iceberg_list_string", "LIST<STRING> → JSON", "JSON array", STATUS],
            ["5.19", "iceberg_list_int", "LIST<INT32> → JSON", "JSON array of ints", STATUS],
            ["5.20", "iceberg_map_string_int", "MAP → JSON", "JSON object", STATUS],
            ["5.21", "iceberg_map_string_string", "MAP → JSON", "JSON object", STATUS],
            ["5.22", "iceberg_nested_struct", "STRUCT{id,inner{x,y}} → JSON", "Nested JSON", STATUS],
            ["5.23", "iceberg_required_string", "STRING (REQUIRED)", "0 nulls", STATUS],
            ["5.24", "iceberg_required_int", "INT (REQUIRED)", "0 nulls", STATUS],
        ],
        col_widths=[1.0, 3.5, 5.5, 4.5, 1.8],
    )

    # ==================================================================
    # TEST 6: Golden Dataset
    # ==================================================================
    doc.add_heading("Test 6: Golden Dataset (All-in-One Smoke Test)", level=1)
    p = doc.add_paragraph()
    p.add_run("Config: ").bold = True
    p.add_run("golden-dataset-extended.json    ")
    p.add_run("Rows: ").bold = True
    p.add_run("1,000 (quick iteration)")

    _add_table(doc,
        ["#", "Validation Query", "Pass Criteria", "Status"],
        [
            ["6.1", "SELECT COUNT(*) FROM golden_dataset_extended", "Returns 1000", STATUS],
            ["6.2", "SELECT * FROM golden_dataset_extended LIMIT 5", "All 36 columns populated with correct types", STATUS],
        ],
        col_widths=[1.0, 12.0, 6.0, 1.8],
    )

    # ==================================================================
    # Execution Checklist
    # ==================================================================
    doc.add_heading("Execution Checklist", level=1)
    _add_table(doc,
        ["Test", "Config", "Table Created", "Ingested", "Queries Pass", "Notes"],
        [
            ["1. Primitive Types", "all-primitive-types.json", "", "", "", ""],
            ["2. Logical Types", "all-logical-types.json", "", "", "", ""],
            ["3. Complex Types", "all-complex-types.json", "", "", "", ""],
            ["4. Null Handling", "null-scenarios.json", "", "", "", ""],
            ["5. Iceberg Coverage", "iceberg-full-coverage.json", "", "", "", ""],
            ["6. Golden Dataset", "golden-dataset-extended.json", "", "", "", ""],
        ],
        col_widths=[3.0, 4.0, 2.5, 2.5, 2.5, 6.0],
    )

    # ==================================================================
    # Known Risks
    # ==================================================================
    doc.add_heading("Known Risks", level=1)
    risks = [
        ("INT96 ingestion",
         "Deprecated, no Parquet column statistics. May fail in some ingestion paths. Not used as timeColumnName."),
        ("BSON as BYTES",
         "Pinot BYTES columns may need special handling in queries (base64 encoding)."),
        ("Timestamp precision loss",
         "MICROS and NANOS timestamps lose sub-millisecond precision in Pinot (stores ms internally)."),
        ("Sentinel collision (Test 4C)",
         "With enableColumnBasedNullHandling: true, sentinel values should be preserved as real data. "
         "If treated as null, column-based null handling is broken."),
        ("Complex types as JSON",
         "All STRUCT/LIST/MAP become JSON strings. JSON_EXTRACT_SCALAR queries work if JSON indexing is enabled."),
        ("DECIMAL as BIG_DECIMAL",
         "Depends on the ParquetToPinotTypeMapper change. If not deployed yet, DECIMAL columns may fail."),
    ]
    for title_text, body_text in risks:
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}: ")
        run.bold = True
        p.add_run(body_text)

    # ---- Save ----
    out_path = "output/Parquet-Type-Testing-Plan.docx"
    doc.save(out_path)
    print(f"Written to: {out_path}")


if __name__ == "__main__":
    build()
